"""
Generate ALTERNATE_DATA_SOURCES_REPORT.docx with McKinsey/Deloitte styling.
Inline formatting only — renders correctly in Google Docs.
"""

import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

# ── Colours (McKinsey/Deloitte palette) ──────────────────────────────────────
NAVY_DARK   = RGBColor(0x1F, 0x38, 0x64)   # #1F3864 — H1
NAVY_MID    = RGBColor(0x2F, 0x54, 0x96)   # #2F5496 — H2
NAVY_LIGHT  = RGBColor(0x1F, 0x49, 0x7D)   # #1F497D — H3
GREY_TEXT   = RGBColor(0x59, 0x59, 0x59)   # body text
GREY_LIGHT  = RGBColor(0xF2, 0xF2, 0xF2)   # table alt row
WHITE       = RGBColor(0xFF, 0xFF, 0xFF)
CALLOUT_BG  = RGBColor(0xEB, 0xF3, 0xFB)   # #EBF3FB — key takeaway box
CALLOUT_TXT = RGBColor(0x1F, 0x38, 0x64)   # navy text in callout
TABLE_HDR   = RGBColor(0x1F, 0x38, 0x64)   # dark navy table header bg
ORANGE_ACC  = RGBColor(0xC5, 0x50, 0x0B)   # figure caption accent
DIVIDER_CLR = RGBColor(0xBF, 0xBF, 0xBF)

BODY_FONT   = "Calibri"
HEAD_FONT   = "Calibri"

# ── XML helpers ───────────────────────────────────────────────────────────────

def rgb_hex(rgb: RGBColor) -> str:
    return f"{rgb[0]:02X}{rgb[1]:02X}{rgb[2]:02X}"


def set_cell_bg(cell, rgb: RGBColor):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    hex_color = rgb_hex(rgb)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def set_cell_borders(cell, top=None, bottom=None, left=None, right=None):
    tc   = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side, val in [("top", top), ("bottom", bottom), ("left", left), ("right", right)]:
        if val:
            el = OxmlElement(f"w:{side}")
            el.set(qn("w:val"),   val.get("val",   "single"))
            el.set(qn("w:sz"),    val.get("sz",    "4"))
            el.set(qn("w:space"), "0")
            el.set(qn("w:color"), val.get("color", "auto"))
            tcBorders.append(el)
    tcPr.append(tcBorders)


def add_para_border(para, side="left", color="1F3864", sz="24", space="144"):
    """Add a left border to a paragraph (callout box effect)."""
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    el = OxmlElement(f"w:{side}")
    el.set(qn("w:val"),   "single")
    el.set(qn("w:sz"),    sz)
    el.set(qn("w:space"), space)
    el.set(qn("w:color"), color)
    pBdr.append(el)
    pPr.append(pBdr)


def add_page_break(doc):
    para = doc.add_paragraph()
    run  = para.add_run()
    br   = OxmlElement("w:br")
    br.set(qn("w:type"), "page")
    run._r.append(br)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after  = Pt(0)


def set_para_shading(para, rgb: RGBColor):
    pPr = para._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    hex_color = rgb_hex(rgb)
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    pPr.append(shd)


def add_horizontal_rule(doc, color="BFBFBF"):
    para = doc.add_paragraph()
    pPr  = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"),   "single")
    bottom.set(qn("w:sz"),    "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    para.paragraph_format.space_before = Pt(4)
    para.paragraph_format.space_after  = Pt(4)


# ── Style helpers ─────────────────────────────────────────────────────────────

def style_run(run, bold=False, italic=False, size=None, color=None, font=BODY_FONT):
    run.font.name  = font
    run.font.bold  = bold
    run.font.italic = italic
    if size:
        run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color


def set_para_spacing(para, before=0, after=6, line=None):
    para.paragraph_format.space_before = Pt(before)
    para.paragraph_format.space_after  = Pt(after)
    if line:
        from docx.shared import Pt as _Pt
        para.paragraph_format.line_spacing = _Pt(line)


# ── Document-level setup ──────────────────────────────────────────────────────

def setup_document():
    doc = Document()
    # Page margins
    for section in doc.sections:
        section.page_width    = Inches(8.5)
        section.page_height   = Inches(11)
        section.left_margin   = Inches(1.0)
        section.right_margin  = Inches(1.0)
        section.top_margin    = Inches(1.0)
        section.bottom_margin = Inches(1.0)
    # Default paragraph style
    style = doc.styles["Normal"]
    style.font.name = BODY_FONT
    style.font.size = Pt(11)
    style.font.color.rgb = GREY_TEXT
    return doc


# ── Heading helpers ───────────────────────────────────────────────────────────

def add_h1(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    style_run(run, bold=True, size=20, color=NAVY_DARK, font=HEAD_FONT)
    set_para_spacing(para, before=18, after=6)
    return para


def add_h2(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    style_run(run, bold=True, size=14, color=NAVY_MID, font=HEAD_FONT)
    set_para_spacing(para, before=14, after=4)
    return para


def add_h3(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    style_run(run, bold=True, size=12, color=NAVY_LIGHT, font=HEAD_FONT)
    set_para_spacing(para, before=10, after=3)
    return para


def add_h2_subtitle(doc, text):
    """Italic subtitle line under an H1."""
    para = doc.add_paragraph()
    run  = para.add_run(text)
    style_run(run, italic=True, size=12, color=NAVY_MID, font=HEAD_FONT)
    set_para_spacing(para, before=0, after=8)
    return para


def add_body(doc, text, bold_spans=None):
    """Add a body paragraph. bold_spans = list of substrings to bold."""
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=6)
    if not bold_spans:
        run = para.add_run(text)
        style_run(run, size=11, color=GREY_TEXT)
    else:
        remaining = text
        for span in bold_spans:
            idx = remaining.find(span)
            if idx > 0:
                r = para.add_run(remaining[:idx])
                style_run(r, size=11, color=GREY_TEXT)
            if idx >= 0:
                r = para.add_run(span)
                style_run(r, bold=True, size=11, color=GREY_TEXT)
                remaining = remaining[idx + len(span):]
        if remaining:
            r = para.add_run(remaining)
            style_run(r, size=11, color=GREY_TEXT)
    return para


def add_callout(doc, text):
    """Key Takeaway blockquote — blue-grey shaded box with left border."""
    para = doc.add_paragraph()
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    set_para_spacing(para, before=8, after=8)
    set_para_shading(para, CALLOUT_BG)
    add_para_border(para, side="left", color="2F5496", sz="18", space="72")
    # Parse bold markers
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if part:
            run = para.add_run(part)
            style_run(run, bold=(i % 2 == 1), size=11, color=CALLOUT_TXT)
    return para


def add_figure_caption(doc, text):
    para = doc.add_paragraph()
    run  = para.add_run(text)
    style_run(run, italic=True, size=10, color=ORANGE_ACC)
    set_para_spacing(para, before=8, after=4)
    return para


def add_code_block(doc, text):
    para = doc.add_paragraph()
    set_para_shading(para, RGBColor(0xF5, 0xF5, 0xF5))
    para.paragraph_format.left_indent  = Inches(0.3)
    para.paragraph_format.right_indent = Inches(0.3)
    set_para_spacing(para, before=6, after=6)
    run = para.add_run(text)
    run.font.name  = "Courier New"
    run.font.size  = Pt(9)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return para


def add_bullet(doc, text):
    para = doc.add_paragraph(style="List Bullet")
    set_para_spacing(para, before=0, after=3)
    para.paragraph_format.left_indent   = Inches(0.3)
    para.paragraph_format.first_line_indent = Inches(-0.2)
    parts = re.split(r"\*\*(.+?)\*\*", text)
    for i, part in enumerate(parts):
        if part:
            run = para.add_run(part)
            style_run(run, bold=(i % 2 == 1), size=11, color=GREY_TEXT)
    return para


def add_checklist_item(doc, text):
    para = doc.add_paragraph()
    set_para_spacing(para, before=0, after=3)
    para.paragraph_format.left_indent = Inches(0.3)
    run = para.add_run("☐  " + text)
    style_run(run, size=11, color=GREY_TEXT)
    return para


# ── Table helpers ─────────────────────────────────────────────────────────────

def add_table(doc, rows_data, header=True, col_widths=None):
    """
    rows_data: list of lists of strings.
    header: first row is header row (navy bg, white text).
    col_widths: list of Inches values.
    """
    if not rows_data:
        return None

    num_cols = max(len(r) for r in rows_data)
    table = doc.add_table(rows=len(rows_data), cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = "Table Grid"

    # Set column widths
    if col_widths:
        for i, col in enumerate(table.columns):
            if i < len(col_widths):
                col.width = col_widths[i]

    for r_idx, row_data in enumerate(rows_data):
        row = table.rows[r_idx]
        row.height = Cm(0.7)
        is_header = (r_idx == 0 and header)
        is_verdict = any(str(cell).strip().lower() == "verdict" for cell in row_data)
        is_alt = (r_idx % 2 == 0 and not is_header)

        for c_idx, cell_text in enumerate(row_data):
            if c_idx >= len(row.cells):
                continue
            cell = row.cells[c_idx]
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

            # Clear default paragraph
            for p in cell.paragraphs:
                p.clear()

            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after  = Pt(3)
            para.paragraph_format.left_indent  = Pt(6)

            # Parse bold markers in cell text
            parts = re.split(r"\*\*(.+?)\*\*", str(cell_text))
            for i, part in enumerate(parts):
                if part:
                    run = para.add_run(part)
                    run.font.name = BODY_FONT
                    run.font.size = Pt(10)
                    if is_header:
                        run.font.bold        = True
                        run.font.color.rgb   = WHITE
                    elif is_verdict:
                        run.font.bold        = (i % 2 == 1)
                        run.font.color.rgb   = NAVY_DARK
                    else:
                        run.font.bold        = (i % 2 == 1)
                        run.font.color.rgb   = GREY_TEXT

            # Background colours
            if is_header:
                set_cell_bg(cell, TABLE_HDR)
            elif is_verdict:
                set_cell_bg(cell, RGBColor(0xEB, 0xF3, 0xFB))
            elif is_alt:
                set_cell_bg(cell, GREY_LIGHT)

    # Light borders
    for row in table.rows:
        for cell in row.cells:
            set_cell_borders(
                cell,
                top    = {"val": "single", "sz": "2", "color": "D9D9D9"},
                bottom = {"val": "single", "sz": "2", "color": "D9D9D9"},
                left   = {"val": "single", "sz": "2", "color": "D9D9D9"},
                right  = {"val": "single", "sz": "2", "color": "D9D9D9"},
            )

    set_para_spacing(doc.add_paragraph(), before=0, after=8)
    return table


# ── Cover page ────────────────────────────────────────────────────────────────

def add_cover_page(doc):
    # Title block
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(48)
    para.paragraph_format.space_after  = Pt(6)
    run = para.add_run("Alternate Data Sources")
    run.font.name  = HEAD_FONT
    run.font.bold  = True
    run.font.size  = Pt(28)
    run.font.color.rgb = NAVY_DARK

    para2 = doc.add_paragraph()
    run2 = para2.add_run("India Health Insurance Platform")
    run2.font.name  = HEAD_FONT
    run2.font.bold  = True
    run2.font.size  = Pt(22)
    run2.font.color.rgb = NAVY_DARK
    para2.paragraph_format.space_after = Pt(12)

    para3 = doc.add_paragraph()
    run3 = para3.add_run("Technical Integration Reference & Vendor Decision Guide")
    run3.font.name   = HEAD_FONT
    run3.font.italic = True
    run3.font.size   = Pt(14)
    run3.font.color.rgb = NAVY_MID
    para3.paragraph_format.space_after = Pt(24)

    add_horizontal_rule(doc)

    doc.add_paragraph()  # spacer

    add_table(doc, [
        ["Field",          "Detail"],
        ["Version",        "1.0"],
        ["Date",           "July 2026"],
        ["Status",         "Living Document — Section 1 complete; Sections 2–21 structured"],
        ["Owner",          "Platform Architecture Team"],
        ["Audience",       "Engineering, Product, Actuarial"],
        ["Classification", "Internal — Confidential"],
        ["Review Cycle",   "Quarterly"],
        ["Next Review",    "October 2026"],
    ], col_widths=[Inches(1.8), Inches(4.8)])

    doc.add_paragraph()

    add_table(doc, [
        ["Version", "Date",      "Author",        "Changes"],
        ["0.1",     "June 2026", "Platform Team", "Initial draft — Environmental section"],
        ["1.0",     "July 2026", "Platform Team", "Full structure; Section 1 complete"],
    ], col_widths=[Inches(0.8), Inches(1.2), Inches(1.8), Inches(2.8)])


# ── Executive Summary ─────────────────────────────────────────────────────────

def add_executive_summary(doc):
    add_page_break(doc)
    add_h1(doc, "Part I — Executive Summary")
    add_horizontal_rule(doc)

    add_h2(doc, "Key Messages")

    add_callout(doc,
        "**1. The platform already has more data than it is using.**\n"
        "iAdore, Karza, and NuralX collectively give us PAN identity, financial profile, biometrics, "
        "and occupation. The gaps are not in coverage — they are in cross-signal contradiction rules. "
        "Implementing Section 20 (internal zero-cost rules) alone will catch 30–40% of fraud cases "
        "before spending a rupee on new vendors."
    )
    add_callout(doc,
        "**2. Three registrations must start today — they take 4–8 weeks and block everything downstream.**\n"
        "IIB FAR (prior rejection history), NHCX (active claims detection), and ABDM sandbox (health records). "
        "None require engineering. All require paperwork. Every week of delay is a week the platform goes live "
        "without India's most important fraud and health signals."
    )
    add_callout(doc,
        "**3. One integration — Account Aggregator — unlocks twelve signals at once.**\n"
        "A single AA consent at Step 2 gives verified income, pharmacy spend (chronic medication proxy), "
        "hospital payments (undisclosed hospitalisation proxy), medical loan EMIs (prior surgery proxy), "
        "alcohol spend, gambling spend, and five more. No other integration has this signal-to-effort ratio."
    )

    add_horizontal_rule(doc)
    add_h2(doc, "What This Document Is")
    add_body(doc,
        "This is the master decision log for every external data source that can enrich underwriting on the "
        "India Health Insurance Platform. For each of the 21 data source categories it records: why the data "
        "is clinically or actuarially valuable, which vendors supply it, what to use now vs. later, exactly "
        "how it connects to the customer journey, and what the regulatory constraints are under IRDAI and DPDPA 2023."
    )
    add_body(doc,
        "Each section is independently actionable. A team can read Section 7 (Diagnostic Lab Data) and "
        "implement it without reading anything else."
    )

    add_horizontal_rule(doc)
    add_h2(doc, "The Five Highest-Priority Decisions")

    # Decision 1
    add_h3(doc, "Decision 1 — IIB FAR + NHCX: Start Registration Today")
    add_body(doc,
        "The Insurance Information Bureau's Fraud Analytics Repository (FAR) catches applicants declined "
        "by another insurer who did not disclose it — the Indian equivalent of the US MIB. NHCX (National "
        "Health Claims Exchange) detects applicants currently hospitalised who are simultaneously applying "
        "for new cover. Neither requires engineering to start. Both require a formal membership process "
        "taking 4–8 weeks.",
        bold_spans=["Fraud Analytics Repository (FAR)", "NHCX"]
    )
    add_table(doc, [
        ["",         "Detail"],
        ["**Action**", "Start IIB membership application and NHCX registration with NHA"],
        ["**Cost**",   "IIB: ₹2–5 lakh/year. NHCX: government pricing"],
        ["**Blocks**", "Section 15 (Fraud Detection) cannot be complete without these"],
    ], col_widths=[Inches(1.4), Inches(5.2)])

    # Decision 2
    add_h3(doc, "Decision 2 — Account Aggregator: Highest ROI Single Integration")
    add_body(doc,
        "One AA consent at Step 2 unlocks verified income, pharmacy spend, hospital payments, medical loan "
        "EMIs, existing insurance detection, alcohol spend, and five more signals. No other integration "
        "returns this breadth from one customer interaction."
    )
    add_table(doc, [
        ["",           "Detail"],
        ["**Vendor**", "Perfios AA (existing iAdore relationship — likely contract extension) or Setu"],
        ["**Cost**",   "₹3–10 per application"],
        ["**Journey Step**", "Step 2 — immediately after PAN verification"],
    ], col_widths=[Inches(1.4), Inches(5.2)])

    # Decision 3
    add_h3(doc, "Decision 3 — Bureau.id: India Fraud Consortium")
    add_body(doc,
        "Bureau.id operates a consortium of 200+ India fintechs and insurtechs sharing device, identity, "
        "and behavioural fraud signals in real time. ICICI Lombard and Bajaj Allianz are already members. "
        "If a device or identity was used in fraud on any member platform, you receive that signal at Step 1."
    )
    add_table(doc, [
        ["",               "Detail"],
        ["**Vendor**",     "Bureau.id (Bangalore)"],
        ["**Cost**",       "₹20–80 per application"],
        ["**Journey Step**", "Step 1 — fires at session start, zero customer friction"],
    ], col_widths=[Inches(1.4), Inches(5.2)])

    # Decision 4
    add_h3(doc, "Decision 4 — ABDM/ABHA: Register on Sandbox Now")
    add_body(doc,
        "ABDM health records are the most powerful health data source in India — actual hospitalisation "
        "history, lab results, prescriptions from NABH/NABL facilities. Production access takes 8–12 weeks "
        "of NHA compliance registration. Sandbox registration takes one day. Start now."
    )
    add_table(doc, [
        ["",           "Detail"],
        ["**Action**", "Register at sandbox.abdm.gov.in today"],
        ["**Cost**",   "Free (government initiative)"],
        ["**Blocks**", "Section 6 (Medical Records) cannot go live without this"],
    ], col_widths=[Inches(1.4), Inches(5.2)])

    # Decision 5
    add_h3(doc, "Decision 5 — CKYC: Verify It Is Already in Your Karza Contract")
    add_body(doc,
        "CKYC (Central KYC Records Registry) lookup is almost certainly already in the Karza contract. "
        "It provides pre-validated identity from another regulated institution and surfaces the mandatory "
        "PEP flag (Politically Exposed Person) required under PMLA. It is a legal compliance requirement."
    )
    add_table(doc, [
        ["",                  "Detail"],
        ["**Action**",        "Call Karza account manager — confirm CKYC is activated"],
        ["**Cost**",          "Likely ₹0 additional (bundled)"],
        ["**Risk of delay**", "PMLA non-compliance if not active before go-live"],
    ], col_widths=[Inches(1.4), Inches(5.2)])

    add_horizontal_rule(doc)
    add_h2(doc, "Implementation Phases at a Glance")
    add_table(doc, [
        ["Phase",        "Timeline",        "What Gets Built",                              "Approx. Cost/Application"],
        ["**Phase 0**",  "Now — 2 weeks",   "Internal fraud rules, district_risk_index DB, DigiLocker extensions, Karza activations, ABDM/IIB registration starts", "Near zero"],
        ["**Phase 1**",  "30–60 days",      "AA, Bureau.id, FingerprintJS, MaxMind, IQAir AQI, lab OCR, CRIF",         "₹5–15"],
        ["**Phase 2**",  "60–180 days",     "ABDM production, BioCatch, HyperVerge, Terra wearables, voice biomarkers, Data Sutram", "₹15–50"],
        ["**Phase 3**",  "6–18 months",     "GVK EMRI, FASTag, pharma direct APIs, mobile app wearables, reinsurer actuarial data",  "Negotiated"],
    ], col_widths=[Inches(0.9), Inches(1.2), Inches(3.5), Inches(1.0)])

    add_horizontal_rule(doc)
    add_h2(doc, "All 21 Data Source Categories — Status Summary")
    add_table(doc, [
        ["#",  "Category",                                      "Status",       "P0 Action",                                          "P1 Action"],
        ["1",  "Environmental & Climate Risk",                  "**Complete**", "—",                                                  "IQAir AQI"],
        ["2",  "Government & Regulatory Data",                  "Placeholder",  "ABDM sandbox, IIB FAR, CKYC",                        "EPFO, AIS/ITR"],
        ["3",  "Financial Behavior Data",                       "Placeholder",  "AA integration",                                     "CRIF HighMark"],
        ["4",  "Identity & KYC Intelligence",                   "Placeholder",  "CKYC, Bureau.id",                                    "HyperVerge"],
        ["5",  "Telecom & Device Intelligence",                 "Placeholder",  "—",                                                  "FingerprintJS, MaxMind"],
        ["6",  "Medical & Health Records",                      "Placeholder",  "ABDM sandbox",                                       "Lab OCR via Karza"],
        ["7",  "Diagnostic Lab Data",                           "Placeholder",  "—",                                                  "Lab OCR (now, free)"],
        ["8",  "Environmental & Geographic Risk (District DB)", "Placeholder",  "Build district_risk_index",                          "Water quality table"],
        ["9",  "Health Infrastructure Data",                    "Placeholder",  "—",                                                  "Data Sutram, MapmyIndia"],
        ["10", "Occupational Risk Data",                        "Placeholder",  "ESIC hazard table",                                  "DGFASLI table"],
        ["11", "Biometric & Advanced Health Signals",           "Placeholder",  "NuralX cross-signal rules",                          "Voice biomarkers (P2)"],
        ["12", "Lifestyle & Behavioral Signals",                "Placeholder",  "—",                                                  "AA pharmacy rules"],
        ["13", "Wearable & Fitness Data",                       "Placeholder",  "—",                                                  "Terra API (P2)"],
        ["14", "Employer & HR Data",                            "Placeholder",  "—",                                                  "GOQii partnership (P2)"],
        ["15", "Fraud Detection Intelligence",                  "Placeholder",  "IIB FAR + NHCX registration",                        "BioCatch, agent rules"],
        ["16", "Term Insurance Specific Signals",               "Placeholder",  "—",                                                  "Declaration questions (now)"],
        ["17", "Actuarial & Insurance Industry Data",           "Placeholder",  "—",                                                  "IRDAI annual data load"],
        ["18", "Geospatial & Satellite Data",                   "Placeholder",  "—",                                                  "Nighttime light table"],
        ["19", "Psychographic & Behavioral Economics",          "Placeholder",  "—",                                                  "Consent velocity JS"],
        ["20", "Internal / Zero-Cost Signals",                  "Placeholder",  "All Section 20 rules",                               "—"],
        ["21", "Emerging & Future Data Sources",                "Placeholder",  "Monitor",                                            "—"],
    ], col_widths=[Inches(0.3), Inches(2.6), Inches(0.85), Inches(1.8), Inches(1.1)])


# ── Part II — Section Template ────────────────────────────────────────────────

def add_section_template(doc):
    add_page_break(doc)
    add_h1(doc, "Part II — Section Template")
    add_horizontal_rule(doc)

    add_body(doc,
        "Every data source section in this document follows the same ten-part structure. "
        "This consistency allows any team member to navigate to any section and find the same "
        "type of information in the same place."
    )
    add_table(doc, [
        ["#",   "Subsection",             "What It Contains"],
        ["1",   "Key Takeaway",           "One blockquoted insight — the single most important thing to know about this data source"],
        ["2",   "Why This Data",          "Clinical or actuarial justification — why this signal improves underwriting decisions"],
        ["3",   "Vendor Analysis",        "Competitive comparison matrix across all relevant vendors — differentiating criteria only"],
        ["4",   "Vendor Selection",       "Clear verdict: what to use now, what to buy later, what to skip entirely"],
        ["5",   "Gaps & Limitations",     "Known weaknesses of the selected vendors — no vendor is perfect"],
        ["6",   "Free Stack vs. Paid",    "Field-by-field comparison where a free alternative exists"],
        ["7",   "Architecture",           "DB lookup vs. live API call — exactly when and where in the 7-step journey"],
        ["8",   "Complete Field Inventory","Every field returned after DB + API calls, with type, source, and UW meaning"],
        ["9",   "Implementation Plan",    "Phased steps with file structure, code sketch, and mock values"],
        ["10",  "Compliance & Regulatory","IRDAI position, DPDPA 2023 requirements, anti-bias rules, migration checklist"],
    ], col_widths=[Inches(0.3), Inches(2.0), Inches(4.3)])

    doc.add_paragraph()
    add_body(doc, "**Data Source Card Format** — used at the top of every individual data source subsection:")
    add_table(doc, [
        ["Attribute",          "Detail"],
        ["**Vendor / Source**","Name and type (government / commercial / free)"],
        ["**Priority**",       "P0 / P1 / P2 / P3 / Monitor"],
        ["**Effort**",         "XS < 1 week | S 1–2 weeks | M 2–6 weeks | L 6–12 weeks | XL 12+ weeks"],
        ["**Estimated Cost**", "Per query or per application"],
        ["**Journey Step**",   "Step 1–7 / STP / POST / STATIC"],
        ["**Signal Strength**","VH / H / M / L / VL"],
        ["**IRDAI Status**",   "Permitted / Mandatory / Caution / Prohibited"],
        ["**DPDPA Category**", "Standard / Sensitive / Biometric / Financial"],
    ], col_widths=[Inches(1.8), Inches(4.8)])


# ── Part III — Baseline ────────────────────────────────────────────────────────

def add_baseline(doc):
    add_page_break(doc)
    add_h1(doc, "Part III — Baseline: Already Integrated")
    add_horizontal_rule(doc)
    add_callout(doc, "**These are live or in active implementation. Do not re-evaluate. Listed here to prevent duplication.**")
    add_table(doc, [
        ["#",    "Source",                "What It Provides",                                                                         "Journey Step",   "API Key"],
        ["B1",   "**iAdore (Perfios)**",  "Demographic profiling: name, DOB, gender, address, employer, occupation, CIBIL score, imputed income, bank statement income, vehicle surrogate income, litigation count, company hazard flag, GST status", "Step 2", "iadore"],
        ["B2",   "**Karza TKYC**",        "PAN verify (name + DOB match), Voter ID verify, Passport verify",                         "Step 2, Step 6", "karza_tkyc"],
        ["B3",   "**Karza OCR Plus**",    "Aadhaar OCR, PAN OCR, bank statement OCR, ITR OCR",                                       "Step 6",         "karza_ocr"],
        ["B4",   "**Karza GSTIN**",       "GST registration status, filing regularity, business name",                               "Step 2",         "karza_gstin"],
        ["B5",   "**Karza VAHAN**",       "Vehicle registration, class, fuel type, insurance status, manufacture year",              "Step 2",         "karza_vahan"],
        ["B6",   "**NuralX (Beaive)**",   "Face vitals via rPPG: heart rate, respiratory rate, BP systolic + diastolic, SpO2, stress index, risk score", "Step 3", "nuralx"],
        ["B7",   "**DigiLocker**",        "Aadhaar XML offline verification + document fetch",                                        "Step 6",         "digilocker"],
        ["B8",   "**Razorpay**",          "Payment processing, Razorpay signature verification",                                     "Payment",        "razorpay"],
        ["B9",   "**Brevo**",             "Transactional email + SMS OTP delivery",                                                  "All steps",      "brevo"],
        ["B10",  "**STP Engine**",        "Automated underwriting decision (APPROVED / REFERRED)",                                    "Step 7",         "stp"],
    ], col_widths=[Inches(0.4), Inches(1.5), Inches(3.0), Inches(1.0), Inches(0.8)])


# ── Section 1 — Environmental & Climate Risk ──────────────────────────────────

def add_section1(doc):
    add_page_break(doc)
    add_h1(doc, "Section 1 — Environmental & Climate Risk")
    add_h2_subtitle(doc, "The one risk factor a customer cannot fabricate")
    add_callout(doc,
        "**Key Takeaway:** A customer in Delhi's Very Poor AQI zone who declares \"no respiratory "
        "conditions\" is statistically implausible — PM2.5 at 180 µg/m³ raises COPD risk by 40% "
        "and IHD risk by 25% independently of any lifestyle factor. This signal costs nothing to "
        "collect and cannot be gamed."
    )
    add_horizontal_rule(doc)

    add_body(doc, "**Section Status**")
    add_table(doc, [
        ["",                      ""],
        ["**Integration Status**","Architecture complete. Free testing stack to be built."],
        ["**Current Stack**",     "OpenWeatherMap + GDACS + Open-Meteo + NASA SEDAC + EM-DAT"],
        ["**Paid Upgrade**",      "Ambee — onboard when first insurer client signs"],
        ["**Journey Steps**",     "Step 2 (DB lookup), Step 3 (live APIs), Step 7 (STP reads stored result)"],
    ], header=False, col_widths=[Inches(1.8), Inches(4.8)])
    add_horizontal_rule(doc)

    # 1.1
    add_h2(doc, "1.1 Why This Data — PM2.5 Is Independent of Self-Declaration")
    add_body(doc,
        "Environmental signals cannot be fabricated. A customer cannot lie about the PM2.5 level "
        "in their district or whether a cyclone hit their pincode last year. This makes geographic "
        "risk a powerful Bayesian prior — a correction factor applied on top of declared health data."
    )
    add_body(doc, "**Clinically validated links (peer-reviewed evidence):**")
    add_table(doc, [
        ["Environmental Signal",               "Health Risk",                         "Magnitude",                                     "Evidence"],
        ["PM2.5 annual mean > 25 µg/m³",       "COPD, IHD, stroke",                   "COPD +15–40%, IHD +10–25%, stroke +10–15%",     "WHO 2021 Global Air Quality Guidelines"],
        ["PM2.5 annual mean > 60 µg/m³",       "Respiratory hospitalisation",          "2× vs clean air zones",                         "ICMR India study 2019"],
        ["Active flood zone (within 100km)",    "Leptospirosis, waterborne disease, dengue clustering", "3–5× disease incidence",       "NCDC India"],
        ["Heat wave days > 30/year",            "Cardiovascular stress, heat stroke mortality", "+12% cardiac mortality",              "IMD / Open-Meteo ERA5"],
        ["High pollen zone",                    "Asthma trigger, allergic rhinitis loading", "+20% asthma exacerbation",               "Ambee India pollen data"],
        ["Industrial zone proximity",           "Heavy metal toxicity, occupational-grade pollution", "Dose-dependent",               "CPCB industrial monitoring"],
    ], col_widths=[Inches(2.0), Inches(1.8), Inches(1.8), Inches(1.0)])

    add_body(doc, "**Permitted insurance use cases** (all approved without additional IRDAI product filing — see Section 1.9):")
    for item in [
        "STP referral trigger (REFERRED only — never auto-REJECTED)",
        "Geographic zone context panel for underwriter reviewer",
        "PED cross-validation flag (high-AQI district + no respiratory PED declared)",
        "Wellness discount eligibility threshold",
        "Mandatory medical exam trigger (high-risk zone + high sum insured)",
    ]:
        add_bullet(doc, item)
    add_horizontal_rule(doc)

    # 1.2
    add_h2(doc, "1.2 Vendor Analysis — Eight Sources Evaluated")
    add_figure_caption(doc, "Figure 1.1 — Vendor Comparison Matrix")
    add_table(doc, [
        ["Criterion",           "Ambee",               "IQAir AirVisual",          "OpenWeatherMap",       "Google Maps AQ",       "CPCB (Govt)",      "GDACS (UN)",     "Open-Meteo",      "NASA SEDAC"],
        ["**India AQI Coverage**","5,000+ cities, pincode-level","Good metros, weak Tier 3","All India, moderate accuracy","All India, CPCB-native","900+ stations only","N/A","N/A","All India 1km grid"],
        ["**Free Tier**",       "15-day trial only",   "10K calls/month",          "1M calls/month",       "10K calls/month",      "Free download",    "Unlimited, free","300K calls/month","Free download"],
        ["**Paid Pricing**",    "₹2–8/query (est.)",   "$29–299/month",            "$40/month",            "$5/1,000 calls",       "N/A",              "N/A",            "Free forever",    "N/A"],
        ["**AQI Standard**",    "US EPA (must convert)","US EPA",                  "US EPA",               "CPCB-native",          "CPCB-native",      "N/A",            "N/A",             "PM2.5 µg/m³"],
        ["**India Pollen Data**","**Yes — only source**","No",                     "No",                   "No",                   "No",               "No",             "No",              "No"],
        ["**Commercial SLA**",  "Yes",                 "Yes",                      "99.5%",                "99.9%",                "None",             "Best-effort",    "Best-effort",     "N/A"],
        ["Verdict",             "Buy at first client", "Skip",                     "**Use now**",          "OWM alternative",      "**Static DB**",    "**Use now**",    "**Use now**",     "**Download once**"],
    ])
    add_horizontal_rule(doc)

    # 1.3
    add_h2(doc, "1.3 Vendor Selection — What to Use and When")
    add_h3(doc, "Now — Free Stack (Before First Insurer Client)")
    add_table(doc, [
        ["API",                         "Purpose",                                         "Free Tier",        "Registration"],
        ["OpenWeatherMap Air Pollution", "Real-time AQI + 6 pollutants",                   "1M calls/month",   "API key — immediate"],
        ["GDACS (UN)",                  "Active disasters within 100km",                   "Unlimited",        "No key needed"],
        ["Open-Meteo",                  "Heat index, dengue season flag, historical ERA5", "300K calls/month", "No key needed"],
        ["NASA SEDAC",                  "20-year PM2.5 annual mean (one-time download)",   "Free",             "Registration required"],
        ["EM-DAT",                      "Historical disaster frequency (one-time download)","Free",            "Academic registration"],
        ["India Post CSV",              "Pincode → lat/lon + district (one-time download)", "Free",            "data.gov.in"],
    ], col_widths=[Inches(2.0), Inches(2.4), Inches(1.2), Inches(1.0)])

    add_h3(doc, "When First Insurer Client Onboards — Add Ambee")
    add_table(doc, [
        ["Ambee API",       "What It Adds Over Free Stack",                          "Necessity"],
        ["Air Quality",     "Better India pincode resolution; commercial SLA",        "Optional — OWM is adequate"],
        ["Natural Disasters","Commercial SLA; better real-time India coverage",       "Optional — GDACS is adequate"],
        ["Pollen",          "**Only India pollen source — no free alternative exists**","**Required for asthma/allergy loading**"],
        ["Weather",         "Unified single-vendor response",                         "Optional — Open-Meteo is sufficient"],
    ], col_widths=[Inches(1.4), Inches(3.0), Inches(2.2)])

    add_h3(doc, "Ambee APIs to Skip Entirely")
    add_table(doc, [
        ["API",              "Reason"],
        ["Wildfire Forecast","North America only — India not covered"],
        ["ILI (Influenza-Like Illness)","US and Europe only"],
        ["Soil",             "Agriculture use case — not relevant to health UW"],
        ["Water Vapor",      "Meteorology use case — not actionable for UW"],
    ], col_widths=[Inches(2.0), Inches(4.6)])
    add_horizontal_rule(doc)

    # 1.4
    add_h2(doc, "1.4 Gaps & Limitations")
    add_callout(doc,
        "**Key gap:** Even after onboarding Ambee, the district_risk_index DB table is still required. "
        "Ambee does not provide historical 20-year PM2.5, district disease burden, or water quality data. "
        "The DB layer and the live API layer are complementary — not substitutes."
    )
    add_table(doc, [
        ["Gap",                                 "Detail",                                          "Impact",                     "Workaround"],
        ["Ambee ILI API does not cover India",  "Influenza surveillance is US/Europe only",        "Cannot use for disease outbreak risk", "NVBDCP data in district_risk_index"],
        ["Ambee Wildfire is North America only","No India wildfire data",                          "—",                          "GDACS active event detection"],
        ["All AQI is US EPA, not CPCB",         "Must convert PM2.5 concentration to CPCB category","Extra computation",        "Conversion table in Section 1.7"],
        ["No historical PM2.5 beyond 1 year",   "Annual mean PM2.5 still needs NASA SEDAC",        "district_risk_index required even with Ambee","NASA SEDAC one-time download"],
        ["No disease burden data (any live API)","No NFHS-5 equivalent available via API",         "district_risk_index required","NFHS-5 district CSV download"],
        ["Ambee free trial only 15 days",       "Cannot test during full platform build",          "—",                          "Free stack for all pre-client testing"],
    ], col_widths=[Inches(1.8), Inches(1.8), Inches(1.4), Inches(1.6)])
    add_horizontal_rule(doc)

    # 1.5
    add_h2(doc, "1.5 Free Stack vs. Ambee — Field-by-Field Comparison")
    add_figure_caption(doc, "Figure 1.2 — Air Quality Field Coverage")
    add_table(doc, [
        ["Field",                       "Ambee",  "OpenWeatherMap (Free)",    "Notes"],
        ["PM2.5 real-time",             "Yes",    "Yes",                      "No gap"],
        ["PM10 real-time",              "Yes",    "Yes",                      "No gap"],
        ["NO2",                         "Yes",    "Yes",                      "No gap"],
        ["NH3 (ammonia)",               "**No**", "**Yes**",                  "OWM is better here"],
        ["CPCB-native AQI",             "No",     "No",                       "Google Maps AQ has CPCB-native; both require conversion"],
        ["Historical 20-year PM2.5",    "**No**", "**No**",                   "NASA SEDAC (free) beats both"],
    ], col_widths=[Inches(2.0), Inches(1.0), Inches(1.8), Inches(1.8)])

    add_figure_caption(doc, "Figure 1.3 — Disaster & Weather Coverage")
    add_table(doc, [
        ["Field",                       "Ambee",   "GDACS (Free)",           "EM-DAT (Free)",             "Open-Meteo (Free)"],
        ["Real-time active events",     "Yes",     "**Yes**",                "No",                        "No"],
        ["Historical 10+ years",        "Limited", "No",                     "**Yes — from 1900**",       "ERA5 from 1940"],
        ["Insurance loss value",        "No",      "No",                     "**Yes**",                   "No"],
        ["Heat wave days historical",   "Limited", "No",                     "No",                        "**ERA5 — best-in-class**"],
        ["Dengue season flag",          "No",      "No",                     "No",                        "**Computed from temp + humidity**"],
    ], col_widths=[Inches(2.0), Inches(1.0), Inches(1.2), Inches(1.4), Inches(1.0)])

    add_figure_caption(doc, "Figure 1.4 — Pollen (Critical Gap)")
    add_table(doc, [
        ["Field",               "Ambee",   "Any Free Alternative"],
        ["Tree pollen (India)", "**Yes**", "**None available**"],
        ["Grass pollen (India)","**Yes**", "**None available**"],
        ["Weed pollen (India)", "**Yes**", "**None available**"],
    ], col_widths=[Inches(2.0), Inches(1.6), Inches(3.0)])
    add_callout(doc,
        "**What you lose without Ambee:** (1) India pollen data — no free alternative exists anywhere. "
        "(2) Commercial SLA — free APIs have no guaranteed uptime. (3) Marginal pincode resolution "
        "improvement for small towns. Everything else is covered adequately by the free stack."
    )
    add_horizontal_rule(doc)

    # 1.6
    add_h2(doc, "1.6 Architecture — Two-Layer Design")
    add_body(doc, "Two complementary layers. Neither replaces the other.")
    add_body(doc, "**Layer 1 — DB Lookup (Static, Loaded Once)**")
    add_body(doc, "Called at Step 2 when pincode is first captured. Zero per-query API cost. Sub-millisecond latency.")
    add_table(doc, [
        ["Table",               "Data Stored",                                                                                  "Source",                                          "Refresh"],
        ["pincode_coords",      "Pincode → lat, lng, district, state",                                                          "India Post CSV (155K pincodes)",                  "Rarely"],
        ["district_risk_index", "PM2.5 annual mean, disaster frequency, heat wave days, disease burden, water quality, composite score", "NASA SEDAC + EM-DAT + Open-Meteo Archive + NFHS-5 + NVBDCP + CPCB", "Annually"],
    ], col_widths=[Inches(1.6), Inches(2.4), Inches(1.8), Inches(0.8)])

    add_body(doc, "**Layer 2 — Live API Calls (Per Application, at Step 3)**")
    add_table(doc, [
        ["API",                         "What It Returns",                          "Cache TTL",           "Latency"],
        ["OpenWeatherMap Air Pollution", "Real-time AQI, PM2.5, NO2, SO2, CO, O3",  "6 hours per pincode", "~200ms"],
        ["GDACS",                       "Active disaster events within 100km",       "6 hours per lat/lng", "~800ms"],
        ["Open-Meteo",                  "Heat index, dengue season flag",            "1 hour per lat/lng",  "~300ms"],
    ], col_widths=[Inches(2.0), Inches(2.4), Inches(1.4), Inches(0.8)])

    add_body(doc, "**Journey Integration Map:**")
    add_table(doc, [
        ["Step",    "Action",                                                                    "Source",   "Latency"],
        ["Step 2",  "Pincode captured → pincode_coords lookup → lat/lng, district stored",       "DB",       "< 1ms"],
        ["Step 2",  "district_risk_index lookup → static risk baseline stored",                  "DB",       "< 1ms"],
        ["Step 3",  "3 live API calls in parallel → environmental snapshot stored on application","Live APIs","~300ms"],
        ["Step 3",  "Compute geographic_risk_score + flags → store on application",              "In-memory","< 1ms"],
        ["Step 7",  "STP engine reads stored score + flags from application record",             "DB read",  "< 1ms"],
    ], col_widths=[Inches(0.7), Inches(3.8), Inches(1.0), Inches(0.7)])
    add_horizontal_rule(doc)

    # 1.7
    add_h2(doc, "1.7 Complete Field Inventory")
    add_figure_caption(doc, "Figure 1.5 — Fields from DB Lookup (Step 2, Static)")
    add_table(doc, [
        ["Field",                   "Type",         "Source",           "Meaning"],
        ["district",                "TEXT",         "pincode_coords",   "District name"],
        ["state",                   "TEXT",         "pincode_coords",   "State name"],
        ["lat",                     "DECIMAL",      "pincode_coords",   "Latitude (for live API calls)"],
        ["lng",                     "DECIMAL",      "pincode_coords",   "Longitude (for live API calls)"],
        ["pm25_annual_mean",        "DECIMAL",      "NASA SEDAC",       "20-year average PM2.5 for district (µg/m³)"],
        ["pm25_zone",               "TEXT",         "Computed",         "Good / Satisfactory / Moderate / Poor / Very Poor / Severe"],
        ["disaster_frequency_score","DECIMAL 0–10", "EM-DAT",           "Historical flood/cyclone/EQ frequency — 20-year district average"],
        ["heat_wave_days_per_year", "INTEGER",      "Open-Meteo Archive","Average annual days above heat wave threshold"],
        ["hypertension_pct",        "DECIMAL",      "NFHS-5",           "% adults with elevated BP in district"],
        ["diabetes_pct",            "DECIMAL",      "NFHS-5",           "% adults with diabetes in district"],
        ["tobacco_use_pct",         "DECIMAL",      "NFHS-5",           "% adults using tobacco in district"],
        ["malaria_cases_per_lakh",  "DECIMAL",      "NVBDCP",           "Annual malaria incidence rate"],
        ["dengue_cases_per_lakh",   "DECIMAL",      "NVBDCP",           "Annual dengue incidence rate"],
        ["fluoride_risk",           "TEXT",         "CPCB",             "none / moderate / high"],
        ["composite_risk_score",    "DECIMAL 0–100","Computed",         "Weighted composite of all district-level fields"],
    ], col_widths=[Inches(1.8), Inches(1.0), Inches(1.4), Inches(2.4)])

    add_figure_caption(doc, "Figure 1.6 — Fields from Live API Calls (Step 3, Current Conditions)")
    add_table(doc, [
        ["Field",                   "Type",    "Source",         "Meaning"],
        ["aqi_current",             "INTEGER", "OpenWeatherMap", "Today's AQI (US EPA scale)"],
        ["aqi_cpcb_category",       "TEXT",    "Computed",       "Good / Satisfactory / Moderate / Poor / Very Poor / Severe"],
        ["pm25_current",            "DECIMAL", "OpenWeatherMap", "Today's PM2.5 concentration (µg/m³)"],
        ["no2_current",             "DECIMAL", "OpenWeatherMap", "Nitrogen dioxide (µg/m³)"],
        ["active_disaster_nearby",  "BOOLEAN", "GDACS",          "Any active disaster event within 100km"],
        ["disaster_type_current",   "TEXT",    "GDACS",          "flood / cyclone / earthquake / wildfire / none"],
        ["heat_index_today",        "DECIMAL", "Open-Meteo",     "Feels-like temperature today (°C)"],
        ["heat_stress_category",    "TEXT",    "Computed",       "Normal / Caution / Danger / Extreme Danger"],
        ["is_dengue_season",        "BOOLEAN", "Computed",       "True if temp + humidity pattern matches dengue season"],
    ], col_widths=[Inches(1.8), Inches(0.9), Inches(1.3), Inches(2.6)])

    add_figure_caption(doc, "Figure 1.7 — Computed Flags Stored on Application")
    add_table(doc, [
        ["Flag",                    "Type",           "Derivation",                                                              "Used For"],
        ["geographic_risk_score",   "DECIMAL 0–100",  "pm25_annual_mean 30% + disaster_frequency 20% + heat_wave_days 15% + aqi_current 20% + active_disaster 15%", "STP referral trigger"],
        ["respiratory_risk_flag",   "BOOLEAN",        "AQI Poor/Very Poor/Severe AND any respiratory PED declared",              "NuralX mandatory + UW referral"],
        ["disaster_risk_flag",      "BOOLEAN",        "disaster_frequency_score > 7 OR active_disaster_nearby = true",           "Sum insured loading trigger"],
        ["heat_mortality_flag",     "BOOLEAN",        "heat_wave_days_per_year > 30 AND heat_stress_category = Danger/Extreme",  "Cardiac risk loading"],
        ["pollen_risk_flag",        "BOOLEAN",        "Ambee pollen above threshold (added when Ambee onboarded)",               "Asthma/allergy loading"],
    ], col_widths=[Inches(1.8), Inches(1.0), Inches(2.8), Inches(1.0)])

    add_figure_caption(doc, "Figure 1.8 — CPCB AQI Conversion (PM2.5 → Indian Standard)")
    add_table(doc, [
        ["PM2.5 Concentration (µg/m³)", "CPCB Category", "UW Action"],
        ["0–30",   "Good",        "No action"],
        ["30–60",  "Satisfactory","No action"],
        ["60–90",  "Moderate",    "No action"],
        ["90–120", "Poor",        "Flag for PED cross-validation"],
        ["120–250","Very Poor",   "Refer to UW if any respiratory PED declared"],
        ["250–500","Severe",      "Refer to UW regardless of declared PED"],
    ], col_widths=[Inches(2.2), Inches(1.6), Inches(2.8)])
    add_horizontal_rule(doc)

    # 1.8
    add_h2(doc, "1.8 Implementation Plan")
    add_h3(doc, "Phase 1 — Free Testing Stack (Now, Before First Insurer Client)")
    add_body(doc, "**Week 1 — One-time DB setup (no API cost):**")
    for item in [
        "Download India Post pincode CSV → load pincode_coords table",
        "Download NASA SEDAC PM2.5 rasters → run Python extraction script → load pm25_annual_mean",
        "Register EM-DAT → download India disaster records → load disaster columns",
        "Pull Open-Meteo ERA5 bulk archive → compute heat wave days → load heat column",
        "Download NFHS-5 district CSVs → load disease burden columns",
        "Download NVBDCP annual report → load malaria, dengue, kala-azar columns",
    ]:
        add_bullet(doc, item)

    add_body(doc, "**Week 1–2 — API client files:**")
    add_code_block(doc,
        "src/lib/external/\n"
        "├── ambee.ts          ← named ambee.ts intentionally; calls OWM now, Ambee later\n"
        "├── disasters.ts      ← GDACS real-time active events\n"
        "├── weather-risk.ts   ← Open-Meteo current heat index + dengue season\n"
        "└── environmental.ts  ← main entry point — calls all 3 in parallel\n\n"
        "src/lib/mock/\n"
        "└── environmental.mock.ts"
    )

    add_body(doc, "**Mock values by pincode prefix (for test mode):**")
    add_table(doc, [
        ["Pincode Prefix", "City",      "Mock AQI", "Mock Category"],
        ["110",            "Delhi",     "180",      "Poor"],
        ["400",            "Mumbai",    "100",      "Moderate"],
        ["560",            "Bengaluru", "55",       "Satisfactory"],
        ["600",            "Chennai",   "70",       "Moderate"],
        ["700",            "Kolkata",   "150",      "Poor"],
        ["default",        "—",         "80",       "Moderate"],
    ], col_widths=[Inches(1.4), Inches(1.4), Inches(1.0), Inches(2.8)])

    add_h3(doc, "Phase 2 — Ambee Upgrade (When First Insurer Client Onboards)")
    add_body(doc, "Three lines change in src/lib/external/environmental.ts. All downstream logic unchanged.")
    add_code_block(doc,
        "// BEFORE (free stack)\n"
        "const [aq, disaster, weather] = await Promise.all([\n"
        "  getAirQuality(pincode),       // OpenWeatherMap\n"
        "  getDisasterRisk(lat, lng),    // GDACS\n"
        "  getWeatherRisk(lat, lng),     // Open-Meteo\n"
        "])\n\n"
        "// AFTER (Ambee)\n"
        "const [aq, disaster, weather, pollen] = await Promise.all([\n"
        "  ambeeGetAirQuality(pincode),    // Ambee — better India resolution\n"
        "  ambeeGetDisasterRisk(lat, lng), // Ambee — commercial SLA\n"
        "  getWeatherRisk(lat, lng),       // Open-Meteo — keep, sufficient\n"
        "  ambeeGetPollen(lat, lng),       // NEW — only India pollen source\n"
        "])"
    )
    add_horizontal_rule(doc)

    # 1.9
    add_h2(doc, "1.9 Compliance & Regulatory")
    add_h3(doc, "IRDAI Position")
    add_body(doc,
        "What requires IRDAI product approval: Explicitly naming AQI or geographic zone as a rated "
        "factor in the product tariff filing."
    )
    add_body(doc, "**What does NOT require approval:**")
    add_table(doc, [
        ["Use Case",                           "Why Permitted"],
        ["STP referral trigger",               "Routing to human UW is operational — REFERRED does not equal higher premium automatically"],
        ["Geographic zone context for UW reviewer","Providing data to a reviewer is not rating"],
        ["PED cross-validation flag",          "Fraud consistency check — not a tariff factor"],
        ["Wellness discount",                  "IRDAI permits wellness discounts (Circular IRDAI/HLT/REG/CIR/194/09/2017)"],
        ["Medical exam threshold",             "UW operational guideline — not a tariff factor"],
    ], col_widths=[Inches(2.2), Inches(4.4)])

    add_callout(doc,
        "**Safe implementation rule:** Geographic risk = REFERRED trigger only, never REJECTED. "
        "Always offer a remedy (NuralX scan, specialist review). Never tell the customer that AQI "
        "is affecting their application."
    )

    add_h3(doc, "Anti-Redlining Rules")
    add_table(doc, [
        ["Rule",                                          "Implementation"],
        ["Never auto-REJECT on geographic score alone",   "geographic_risk_score triggers REFERRED only"],
        ["Always offer remedy",                           "If respiratory_risk_flag = true, offer free NuralX scan"],
        ["Never surface AQI to customer",                 "No geographic AQI shown on any customer-facing screen"],
        ["Cap geographic contribution",                   "Geographic score ≤ 30% weight in overall STP composite score"],
        ["Quarterly monitoring",                          "Flag any district with > 2× platform average referral rate"],
    ], col_widths=[Inches(2.8), Inches(3.8)])

    add_h3(doc, "DPDPA 2023 Requirements")
    add_table(doc, [
        ["Requirement",                           "Action Required"],
        ["Data Processing Agreement",             "Required with OpenWeatherMap, GDACS, Open-Meteo, Ambee before go-live"],
        ["Section 16 automated decision disclosure","Inform customer at Step 1 consent screen that automated processing is used"],
        ["Data retention",                        "Environmental snapshot follows application retention (policy + 7 years)"],
        ["Audit trail",                           "All API calls logged via callExternalAPI() to api_call_logs table"],
    ], col_widths=[Inches(2.4), Inches(4.2)])

    add_h3(doc, "Quarterly Monitoring Checklist")
    for item in [
        "Referral rate by district — flag any district > 2× platform average",
        "Referral rate by state — flag any state > 2× platform average",
        "Verify no REJECTED applications where rejection is primarily geographic",
        "Check if NASA SEDAC released new annual PM2.5 rasters — refresh district_risk_index if yes",
    ]:
        add_checklist_item(doc, item)
    add_horizontal_rule(doc)

    # 1.10
    add_h2(doc, "1.10 Migration Path — Free Stack to Ambee")
    add_h3(doc, "Ambee Onboarding Checklist")
    for item in [
        "Sign Ambee contract + Data Processing Agreement",
        "Receive API key → add to .env.local and Vercel environment as AMBEE_API_KEY",
        "Implement ambeeGetAirQuality() in src/lib/external/ambee.ts",
        "Implement ambeeGetDisasterRisk() in src/lib/external/ambee.ts",
        "Implement ambeeGetPollen() in src/lib/external/ambee.ts",
        "Add pollen_risk_flag field to application schema and STP payload",
        "Update src/lib/external/environmental.ts — 3-line change",
        "Test in test mode (mock) → test in live mode with real Ambee key",
        "Toggle insurer to live mode via Admin portal",
    ]:
        add_checklist_item(doc, item)

    add_h3(doc, "Free Stack Registration Reference")
    add_table(doc, [
        ["API",             "Registration URL",                   "Free Tier"],
        ["OpenWeatherMap",  "openweathermap.org/api",             "1M calls/month"],
        ["GDACS",           "gdacs.org/gdacsapi",                 "Free, no key required"],
        ["Open-Meteo",      "open-meteo.com",                     "300K calls/month, no key required"],
        ["NASA SEDAC",      "sedac.ciesin.columbia.edu",          "Free download, registration required"],
        ["EM-DAT",          "emdat.be",                           "Free, academic registration"],
        ["NFHS-5",          "rchiips.org/nfhs/nfhs5.shtml",       "Free download"],
        ["India Post Pincodes","data.gov.in",                     "Free download"],
    ], col_widths=[Inches(1.4), Inches(2.8), Inches(2.4)])


# ── Placeholder section helper ────────────────────────────────────────────────

def add_placeholder_section(doc, section_num, title, subtitle, key_takeaway, status_rows, subsections):
    add_page_break(doc)
    add_h1(doc, f"Section {section_num} — {title}")
    add_h2_subtitle(doc, subtitle)
    add_callout(doc, f"**Key Takeaway:** {key_takeaway}")
    add_horizontal_rule(doc)
    add_body(doc, "**Section Status**")
    add_table(doc, [["", ""]] + status_rows, header=False, col_widths=[Inches(1.8), Inches(4.8)])
    add_horizontal_rule(doc)
    for ss in subsections:
        add_h3(doc, ss)
        para = doc.add_paragraph()
        run = para.add_run("Content to be completed.")
        run.font.name   = BODY_FONT
        run.font.size   = Pt(10)
        run.font.italic = True
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        set_para_spacing(para, before=0, after=8)


# ── Appendix helper ───────────────────────────────────────────────────────────

def add_appendix(doc, letter, title, description, subsections):
    add_page_break(doc)
    add_h1(doc, f"Appendix {letter} — {title}")
    add_horizontal_rule(doc)
    para = doc.add_paragraph()
    run  = para.add_run(description)
    run.font.name   = BODY_FONT
    run.font.size   = Pt(11)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    set_para_spacing(para, before=0, after=12)
    for ss in subsections:
        add_h3(doc, ss)
        p = doc.add_paragraph()
        r = p.add_run("Content to be completed.")
        r.font.name   = BODY_FONT
        r.font.size   = Pt(10)
        r.font.italic = True
        r.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        set_para_spacing(p, before=0, after=8)


# ── Main ──────────────────────────────────────────────────────────────────────

def main():
    doc = setup_document()

    add_cover_page(doc)
    add_executive_summary(doc)
    add_section_template(doc)
    add_baseline(doc)

    # Part IV header
    add_page_break(doc)
    add_h1(doc, "Part IV — Data Source Sections")
    add_horizontal_rule(doc)

    add_section1(doc)

    # Sections 2–21 — structured placeholders
    placeholders = [
        (2, "Government & Regulatory Data",
         "The highest-quality signals in India are free — and require paperwork, not engineering",
         "ABDM health records, AA financial data, CKYC identity, and AIS income verification are all government-backed, free or near-free, and consent-based. The barrier is not technical — it is registration paperwork that takes 4–12 weeks.",
         [["**Integration Status**","Placeholder — to be completed"],["**P0 Actions**","ABDM sandbox registration, IIB FAR registration, CKYC activation in Karza"],["**Key Vendors**","NHA (ABDM), CERSAI (CKYC), CBDT (AIS/ITR), EPFO, UIDAI, ESIC"]],
         ["2.1 Why This Data — Government Sources Are Ground Truth","2.2 ABDM / ABHA — Ayushman Bharat Digital Mission Health Records","2.3 Account Aggregator (AA) Framework","2.4 EPFO — Employee Provident Fund Organisation","2.5 ESIC — Employee State Insurance Corporation","2.6 CKYC Registry — Central KYC Records Registry","2.7 AIS / Form 26AS / ITR — Income Tax Records","2.8 VAHAN / SARATHI — Enhanced Signals Beyond iAdore","2.9 MCA21 — Company Registry Deep Signals","2.10 RBI Wilful Defaulter List","2.11 SEBI Debarred Entities List","2.12 UIDAI Aadhaar e-KYC (OTP-Based + Death Detection)","2.13 NCRB — National Crime Records Bureau","2.14 PM-JAY / Ayushman Bharat Card (BPL Income Indicator)","2.15 Architecture — Journey Integration Map","2.16 Compliance & Regulatory Notes"]),
        (3, "Financial Behavior Data",
         "What people pay for reveals what they won't declare",
         "Regular pharmacy UPI payments to 1mg or PharmEasy prove chronic medication use more reliably than any health declaration form. A customer paying ₹800/month at Apollo Pharmacy for the last 8 months is almost certainly managing a chronic condition — regardless of what they wrote on Step 4.",
         [["**Integration Status**","Placeholder — to be completed"],["**P0 Action**","Account Aggregator integration — highest ROI single integration"],["**Key Vendors**","Perfios AA, Setu, CRIF HighMark, Experian India, Finvu, Onemoney"]],
         ["3.1 Why This Data — Spending Reveals Health Behaviour","3.2 Account Aggregator — The Single Consent That Unlocks Twelve Signals","3.3 UPI Transaction Pattern Analysis — Merchant-Category Rules","3.4 Pharmacy Spend Extraction Logic","3.5 Hospital & Medical Payment Detection","3.6 Medical Loan EMI Detection (Prior Surgery Proxy)","3.7 CRIF HighMark — Multi-Bureau (MFI + Medical Loan Signals)","3.8 NACH / eMandate Active Mandates (via AA)","3.9 Experian India — Second Bureau View","3.10 Architecture — Journey Integration Map","3.11 Compliance & Regulatory Notes"]),
        (4, "Identity & KYC Intelligence",
         "Know who is applying before they tell you anything about their health",
         "Bureau.id's consortium of 200+ India platforms means that if a device or mobile number was used in fraud anywhere in the India fintech ecosystem, you know it at Step 1 — before the customer has typed a single character of health information.",
         [["**Integration Status**","Placeholder — to be completed"],["**P0 Actions**","Activate CKYC in Karza; Bureau.id integration"],["**Key Vendors**","Bureau.id, HyperVerge, Gridlines, Signzy, Setu, IDfy"]],
         ["4.1 Why This Data — Identity Fraud Is the Gateway to All Other Fraud","4.2 Bureau.id — India Fraud Intelligence Consortium","4.3 HyperVerge — AI-Native Face Match and Liveness Detection","4.4 Gridlines — Court Records and Identity Verification","4.5 Signzy — Video KYC (IRDAI V-CIP Approved)","4.6 Architecture — Journey Integration Map","4.7 Compliance & Regulatory Notes (PMLA, DPDPA, IRDAI V-CIP Circular)"]),
        (5, "Telecom & Device Intelligence",
         "The device tells you what the customer won't",
         "A FingerprintJS device ID that has submitted 14 applications with 14 different PAN numbers in the last 30 days is a fraud ring — identifiable at ₹0.001 per session. This is the cheapest signal in the entire document.",
         [["**Integration Status**","Placeholder — to be completed"],["**P1 Actions**","FingerprintJS Pro (npm install — 2 hours), MaxMind IP (REST API — 1 day)"],["**Key Vendors**","FingerprintJS Pro, MaxMind GeoIP2, BioCatch, Truecaller Enterprise, NumVerify"]],
         ["5.1 Why This Data — Session Signals Are Unforgeable","5.2 FingerprintJS Pro — Stable Device ID Across Incognito and Browser Restarts","5.3 MaxMind GeoIP2 — IP Geolocation and VPN/Proxy/TOR Detection","5.4 BioCatch — Behavioural Biometrics","5.5 Truecaller Enterprise — Mobile Number Intelligence","5.6 Architecture — Journey Integration Map","5.7 Compliance & Regulatory Notes (DPDPA Section 7d — Fraud Prevention Exemption)"]),
        (6, "Medical & Health Records",
         "The ABHA consent gate is 60 seconds of customer friction for ground-truth PED verification",
         "When a customer shares ABHA records, their prescription history is visible. If they are on Metformin, they have diabetes — regardless of what they declared. The drug-to-condition mapping table could detect 60–70% of PED misrepresentation cases once ABDM is live.",
         [["**Integration Status**","Placeholder — to be completed"],["**P0 Action**","Register on ABDM sandbox today — production takes 8–12 weeks"],["**Key Vendors**","NHA (ABDM), Eka.care, Practo, Apollo 24x7 — all via ABDM path"]],
         ["6.1 Why This Data — Self-Declaration Is the Weakest Link","6.2 ABDM — Full Health Record Access via HIE-CM (FHIR R4)","6.3 ABHA Prescription Analysis — Drug-to-Condition Mapping Table","6.4 Psychotropic Medication Flag (Most Underreported Condition in India)","6.5 PHR Aggregators — Eka.care, Practo, Apollo 24x7","6.6 CoWIN — Vaccination Records (via DigiLocker — Already Integrated)","6.7 GVK EMRI — Ambulance Call Records","6.8 Architecture — ABDM Consent Flow and Journey Integration","6.9 Compliance & Regulatory Notes (IRDAI Circular 123/2023, DPDPA)"]),
        (7, "Diagnostic Lab Data",
         "A lab report uploaded today costs nothing extra — and HbA1c does not lie",
         "Adding an optional lab report upload field at Step 4 with Karza OCR costs zero — Karza is already integrated, OCR is already contracted. HbA1c > 6.5% when the customer declared 'no diabetes' is an automatic contradiction flag with no ambiguity.",
         [["**Integration Status**","Placeholder — to be completed"],["**Immediate Win**","Lab report OCR at Step 4 — zero new vendor, zero new cost, implement this week"],["**Key Vendors**","Thyrocare, Dr Lal PathLabs, SRL, Metropolis"]],
         ["7.1 Why This Data — Objective Biomarkers Cannot Be Fabricated","7.2 Near-Term Approach — Lab Report OCR via Karza (Implement Now, Zero Cost)","7.3 Lab Marker Interpretation Table (HbA1c, LDL, TSH, eGFR, Cotinine, ALT, Creatinine)","7.4 NABL Accreditation Verification","7.5 Thyrocare / Dr Lal PathLabs — Direct API (Long-Term via ABDM)","7.6 Architecture — Journey Integration Map","7.7 Compliance & Regulatory Notes"]),
        (8, "Environmental & Geographic Risk (District DB Layer)",
         "A one-week DB build delivers a lifetime of free district-level risk intelligence",
         "The district_risk_index table costs one week to build, nothing to query, and never expires meaningfully. A customer from Muzaffarpur (Bihar's kala-azar endemic belt) who declares 'no tropical disease history' is statistically worth a second look — and this knowledge costs ₹0 per application.",
         [["**Integration Status**","Placeholder — to be completed"],["**Phase 0 Action**","Build district_risk_index — free government data, 1 week effort"],["**Key Sources**","NFHS-5, NVBDCP, NCDC, CPCB, NASA SEDAC, EM-DAT, NDMA, IMD/Open-Meteo"]],
         ["8.1 Why This Data — Static District Intelligence Is the Foundation of All Geographic Risk","8.2 district_risk_index — Table Schema and Column Definitions","8.3 NFHS-5 — Disease Burden by District","8.4 IHME GBD — DALYs by State/District","8.5 NVBDCP — Vector Disease Endemic Zones","8.6 NCDC — TB and Kala-Azar Incidence by District","8.7 CPCB — Water Quality Contamination Zone Map","8.8 EM-DAT — Historical Disaster Frequency and Insurance Loss by District","8.9 Open-Meteo ERA5 Archive — Annual Heat Wave Days by District","8.10 NASA VIIRS — Nighttime Light Intensity (Area Wealth Proxy)","8.11 Architecture — One-Time Load + Annual Refresh Process","8.12 Implementation Plan — Python Extraction Scripts and Postgres Load"]),
        (9, "Health Infrastructure Data",
         "Drive time to the nearest ICU is a claim severity predictor, not just a coverage metric",
         "A high-risk cardiac applicant in a district where the nearest cath lab is 4 hours away has objectively worse prognosis for a cardiac event than the same applicant 20 minutes from AIIMS. This is claim severity data, not just access quality — and it feeds pricing models.",
         [["**Integration Status**","Placeholder — to be completed"],["**Key Vendors**","NABH, NHP, Data Sutram, MapmyIndia, Google Maps Distance Matrix"]],
         ["9.1 Why This Data — Hospital Access Affects Claim Severity","9.2 NABH Hospital Database (2,500+ Accredited Hospitals — Free Download)","9.3 NHP — National Health Portal Hospital Registry (Free API)","9.4 ICU and Cardiac Cath Lab Drive Time (Google Maps Distance Matrix)","9.5 PMJAY Empanelled Hospital List","9.6 MapmyIndia — India-Native Hospital and Specialist POI","9.7 Architecture — Static DB Table: pincode_health_infrastructure","9.8 Compliance & Regulatory Notes"]),
        (10, "Occupational Risk Data",
         "iAdore gives a binary flag; a 4-tier classification gives a loading multiplier",
         "The difference between 'isHazardous: true' (binary) and 'occupational_risk_tier: 3, loading_multiplier: 1.50' (quantitative) is the difference between a flag that triggers manual UW and a number that feeds directly into the STP score. The 4-tier ESIC classification is a free, one-day build.",
         [["**Integration Status**","Placeholder — to be completed"],["**Phase 0 Action**","Build ESIC 4-tier hazard classification lookup table — free, 1 day"],["**Key Sources**","ESIC Schedule, DGFASLI, iAdore (existing), VAHAN/SARATHI (existing)"]],
         ["10.1 Why This Data — Occupation Is One of the Three Core UW Variables","10.2 4-Tier Hazard Classification — ESIC + DGFASLI (Loading Multipliers)","10.3 DGFASLI — Occupational Fatality Rate Table by NIC Code","10.4 VAHAN / SARATHI — Commercial Driver and HMV License Flags","10.5 DG Shipping — Merchant Navy / Seafarer CDC Verification","10.6 DGCA — Commercial Pilot and Aviation Crew Confirmation","10.7 ESIC Card — Factory Worker + Income Band Confirmation","10.8 Architecture — Journey Integration Map","10.9 Compliance & Regulatory Notes"]),
        (11, "Biometric & Advanced Health Signals",
         "NuralX already captures the data — the cross-signal rules are what is missing",
         "A NuralX systolic BP reading of 155 mmHg from a customer who declared 'no hypertension' is not a flag — it is a contradiction. The contradiction rules cost nothing to implement and make the existing NuralX integration significantly more powerful today.",
         [["**Integration Status**","Placeholder — to be completed. NuralX already integrated."],["**Immediate Win**","Add NuralX cross-signal contradiction rules — zero new vendor, zero cost"],["**Key Vendors**","NuralX (existing), Vocalis Health, Lapetus Solutions, Nuralogix"]],
         ["11.1 NuralX — Cross-Signal Contradiction Rules to Add Now (Zero Cost)","11.2 Nuralogix Anura — Alternative Face Vitals (Biological Age Signal)","11.3 Voice Biomarkers — COPD, Diabetes, and Depression Detection","11.4 Biological Age Estimation from Selfie (Lapetus)","11.5 Pulse Wave Analysis — Arterial Stiffness","11.6 Architecture — Journey Integration Map","11.7 Compliance & Regulatory Notes (Biometric Data — DPDPA Sensitive Category)"]),
        (12, "Lifestyle & Behavioral Signals",
         "The fastest path to lifestyle signals is already inside the AA bank statement",
         "The Account Aggregator integration already captures pharmacy spend, alcohol delivery payments, and gym subscriptions from bank transaction data. No separate lifestyle API is needed. Direct partnerships with food delivery or ride-sharing platforms are not worth pursuing.",
         [["**Integration Status**","Placeholder — to be completed"],["**Note**","Most signals in this section flow through AA — no separate vendor needed once AA is live"]],
         ["12.1 Why This Data — Lifestyle Signals Contradict or Confirm Self-Declaration","12.2 Pharmacy Purchase History — AA Path vs Direct Partnership (AA Wins)","12.3 OPD Consultation History — ABDM Path vs Direct Partnerships (ABDM Wins)","12.4 Fitness App Data — GOQii Enterprise, Strava","12.5 Mental Health App Usage — AA Subscription Detection","12.6 Signals to Skip — Food Delivery and Ride-Sharing","12.7 Architecture — Journey Integration Map","12.8 Compliance & Regulatory Notes"]),
        (13, "Wearable & Fitness Data",
         "10% of applicants own a compatible wearable — position it as a discount, not a requirement",
         "Resting heart rate averaged over 7 days from an Apple Watch is a stronger cardiac health signal than a single 30-second NuralX scan. Terra API connects 50+ wearable platforms in one SDK. But India penetration is ~10% — make it opt-in, discount-rewarding, and never penalising.",
         [["**Integration Status**","Placeholder — to be completed"],["**Key Vendor**","Terra API — single SDK for 50+ wearable platforms"]],
         ["13.1 Why This Data — Longitudinal Wearable Data Beats Point-in-Time Scans","13.2 India Coverage Reality (~10%) and Opt-In Discount Strategy","13.3 Terra API — 50+ Platforms in One SDK","13.4 Key Wearable Signals — Resting HR, HRV, VO2max, Sleep SpO2, CGM","13.5 Vital API — Alternative Aggregator","13.6 Architecture — Journey Integration Map","13.7 IRDAI Posture and Compliance Notes (Discount-Not-Penalty Framing)"]),
        (14, "Employer & HR Data",
         "The most common health insurance fraud vector in India is the ex-employee hiding a group health claim",
         "An employee laid off after a diabetes diagnosis claimed under corporate group health cover — who now applies individually and declares 'no diabetes' — is the single most frequent fraud pattern in Indian health insurance. IIB FAR catches this. The fix is in Section 15, not here.",
         [["**Integration Status**","Placeholder — to be completed"],["**Key Vendors**","IIB (group claims via Section 15), GOQii Enterprise, Razorpay Payroll"]],
         ["14.1 Why This Data — Group-to-Individual Policy Transition Is the Highest-Frequency Fraud Vector","14.2 Group Health Insurance Claims History — Detection via IIB FAR (See Section 15)","14.3 GOQii Enterprise Health Score — Longitudinal Wellness Trend","14.4 Payroll Platform Data — Razorpay Payroll, greytHR, Darwinbox, Keka","14.5 Background Verification for High-Value Applications","14.6 Architecture — Journey Integration Map","14.7 Compliance & Regulatory Notes"]),
        (15, "Fraud Detection Intelligence",
         "IIB FAR and NHCX are the two most important integrations not yet on the roadmap",
         "IIB FAR tells you if this applicant was rejected by another insurer. NHCX tells you if this applicant is currently hospitalised. Both are mandatory compliance registrations for a serious health insurer. Neither requires engineering to start — only paperwork. Start both today.",
         [["**Integration Status**","Placeholder — to be completed"],["**P0 Actions**","IIB membership application + IIB FAR agreement + NHCX registration with NHA — start today"],["**Key Vendors**","IIB, NHCX/NHA, Bureau.id (Section 4), BioCatch (Section 5)"]],
         ["15.1 Why This Data — Three IIB Products, Not One","15.2 IIB Basic Claims Lookup (Prior Claims History Across Member Insurers)","15.3 IIB FAR — Fraud Analytics Repository (Prior Rejection / Postponement History)","15.4 NHCX — National Health Claims Exchange (Active Claims Detection in Real Time)","15.5 Hospital Blacklist Database — Build Internally from Free Sources","15.6 Agent / Broker Fraud Detection Rules — Internal STP Engine Rules","15.7 Cluster / Syndicate Fraud Detection Rules — Batch Analysis","15.8 Statistical Anomaly Flags — PED Underreporting Rate vs NFHS-5 District Prevalence","15.9 Architecture — Journey Integration Map","15.10 Compliance & Regulatory Notes (IRDAI Circular IRDA/SDD/CIR/019/2011)"]),
        (16, "Term Insurance Specific Signals",
         "Three declaration questions added to the form today cost nothing and are legally required anyway",
         "'Have you traveled to a conflict zone in the last 12 months?', 'Do you participate in adventure sports?', 'Are you currently serving in the armed forces?' — these are standard underwriting questions for term insurance. They cost nothing to add and create the legal basis for exclusions. Add them in the next sprint.",
         [["**Integration Status**","Placeholder — to be completed"],["**Immediate Actions**","Add 3 declaration questions to health form — zero cost, zero new vendor"]],
         ["16.1 Why This Data — Term Insurance Has a Different Risk Profile to Health","16.2 High-Risk Occupation Table — Mortality Multipliers by Occupation","16.3 Zero-Cost Declaration Questions — Foreign Travel, Adventure Sports, Military Service","16.4 Smoking Cross-Verification — AA UPI + NuralX + Cotinine Test","16.5 Alcohol Consumption Cross-Verification — AA UPI Alcohol Delivery Detection","16.6 PMJJBY / PMSBY Claims History — via IIB","16.7 Architecture — Journey Integration Map","16.8 Compliance & Regulatory Notes (IRDAI Suicide Mandate, DPDPA Mental Health Sensitivity)"]),
        (17, "Actuarial & Insurance Industry Data",
         "IRDAI publishes India's actual claims experience every year — for free",
         "The IRDAI annual report contains disease-wise claims frequency, geographic claims distribution, and fraud rejection statistics for the entire Indian health insurance market. Loading this annually into the STP scoring calibration is a free, half-day task. Not doing it means the STP model is not using the most current India claims experience.",
         [["**Integration Status**","Placeholder — to be completed"],["**Key Sources**","IRDAI Annual Report (free), IIB, GIC Re, Munich Re, Swiss Re, TPA networks"]],
         ["17.1 Why This Data — Actuarial Tables Must Reflect Actual India Experience","17.2 IRDAI Annual Report — Disease Statistics, Claims Frequency, Fraud Data (Free, Annual)","17.3 GIC Re — India Mortality and Morbidity Tables (via Reinsurance Partnership)","17.4 Munich Re / Swiss Re India — Proprietary Actuarial Intelligence","17.5 TPA Claims Data — Medi Assist, Raksha, Health India, Paramount","17.6 IHME Global Burden of Disease — Commercial License for District-Level DALYs","17.7 Architecture — Annual Calibration Process","17.8 Compliance & Regulatory Notes"]),
        (18, "Geospatial & Satellite Data",
         "Satellite data from NASA is free, globally available, and more accurate than self-declared income for rural applicants",
         "NASA VIIRS nighttime light intensity at a pincode level is a robust income proxy for rural and semi-urban applicants where no formal financial footprint exists. A farmer in a dark district declaring ₹15 lakh annual income is worth a second look. The data is free and the table builds in one day.",
         [["**Integration Status**","Placeholder — to be completed"],["**Note**","Most signals here feed into district_risk_index (Section 8) — not standalone integrations"]],
         ["18.1 Why This Data — Satellite Signals Are Objective and Cannot Be Fabricated","18.2 Nighttime Light Intensity — NASA VIIRS / NOAA (Area Wealth Proxy, Free)","18.3 Satellite-Based Crop Yield — ISRO Bhuvan / Sentinel-2 (Farmer Income Verification)","18.4 Flood Inundation Maps — ISRO Bhuvan (feeds district_risk_index)","18.5 Urban Sprawl / NDVI Change — Sentinel-2 Copernicus","18.6 Road Accident Hotspot — iRAD / NCRB","18.7 Architecture and Implementation Notes","18.8 Compliance & Regulatory Notes"]),
        (19, "Psychographic & Behavioral Economics",
         "The application form itself generates fraud signals — at zero cost",
         "A customer who accepts all Step 1 consent checkboxes in 2.3 seconds did not read them. A customer who copy-pastes answers into every health declaration field submitted a pre-prepared template. These signals exist in every session, cost nothing to collect, and require only JavaScript timestamps to capture.",
         [["**Integration Status**","Placeholder — to be completed"],["**Note**","All signals in this section are internal and zero-cost. Implement in Phase 0."]],
         ["19.1 Why This Data — Behaviour During Application Is a Fraud Signal","19.2 Price Sensitivity at Plan Selection — Engagement Pattern Analytics (Step 4)","19.3 Application Abandonment Patterns — Which Step and Which Field","19.4 Consent Acceptance Velocity — Checkbox Interaction Timestamps","19.5 Investment Behaviour / Risk Appetite — via AA Demat Data","19.6 Architecture — JavaScript Tracking Implementation","19.7 Compliance & Regulatory Notes (DPDPA — Behavioural Data as Personal Data)"]),
        (20, "Internal / Zero-Cost Signals",
         "The most valuable fraud signals on this platform cost nothing — they use data already collected",
         "Declaring 'no diabetes' while NuralX returns BP 158/96, BMI is 32, age is 52, and the AA bank statement shows ₹1,200/month at Apollo Pharmacy is not a declaration — it is a case. All four signals are already collected. Only the contradiction rule is missing.",
         [["**Integration Status**","Placeholder — to be completed"],["**Phase 0 Action**","Implement all rules in this section. Zero external API. Zero cost. Target: 2 weeks."]],
         ["20.1 Why This Data — Cross-Signal Contradiction Requires No New Vendors","20.2 Session Behavioural Signals — Copy-Paste, Field Dwell Time, Form Completion Speed, GPS vs Pincode","20.3 Application-Level Cross-Reference Rules — Same PAN, Same Device, Same IP, Same Doctor","20.4 NuralX Contradiction Rules — BP vs Hypertension Denial, HR vs Smoker Denial, SpO2 vs Respiratory Denial","20.5 AA Contradiction Rules — Pharmacy Spend vs No PED, Hospital Payment vs No Hospitalisation","20.6 Statistical Anomaly Flags — Batch Analysis: District PED Rate vs NFHS-5, Agent Claim Rate","20.7 Nominee Anomaly Flags — Age Gap, Non-Family Nominee, Nominee Aadhaar Status","20.8 Architecture — STP Engine Rules and Batch Analysis Schedule","20.9 Implementation Plan — Phase 0, Zero Cost"]),
        (21, "Emerging & Future Data Sources",
         "Monitor these — do not build yet",
         "Genomic data is currently prohibited for insurance use under proposed DPDPA regulations. Telco-AA is not yet live. ONDC has no insurance B2B API. The right action for everything in this section is to assign one person to monitor regulatory developments quarterly — not to build.",
         [["**Integration Status**","Monitor only — not actionable in Phase 1"]],
         ["21.1 ONDC — Open Network for Digital Commerce (Pharmacy + Health Purchase Data, 2027+)","21.2 Telco-AA — Jio / Airtel as Financial Information Providers (Sandbox Expected 2025–26)","21.3 DEPA — Health AA and Telecom AA Extension (NITI Aayog Roadmap, 2026+)","21.4 DigiYatra — Airport Biometric Travel Frequency (No B2B API Currently)","21.5 Genomics / DNA Health Risk — Currently Prohibited Under Proposed DPDPA","21.6 Mental Health Platform Data — Pending IRDAI Mandatory Coverage Regulatory Outcome","21.7 OCEN — Open Credit Enablement Network (MSME Business Health, P3)","21.8 Monitoring Cadence and Review Triggers"]),
    ]

    for args in placeholders:
        add_placeholder_section(doc, *args)

    # Appendices
    add_appendix(doc, "A", "Vendor Master Directory",
        "To be completed — full vendor contact, pricing, contract notes, and evaluation status.",
        ["A.1 Already Integrated Vendors","A.2 P0 Vendors — Start Immediately","A.3 P1 Vendors — 30–60 Days","A.4 P2 Vendors — 60–180 Days","A.5 P3 Vendors — 6–18 Months","A.6 Vendors to Evaluate","A.7 Vendors to Skip — with Reasons"])

    add_appendix(doc, "B", "Priority Implementation Roadmap",
        "To be completed — full phased action plan with effort, cost, dependencies, and owner.",
        ["B.1 Phase 0 — Immediate (0–2 Weeks, Zero/Near-Zero Cost)","B.2 Phase 1 — 30–60 Days (High ROI, Clear Path)","B.3 Phase 2 — 60–180 Days (Medium Complexity, High Value)","B.4 Phase 3 — 6–18 Months (Strategic / Partnerships / Regulatory Approval)","B.5 Dependency Map — What Unlocks What"])

    add_appendix(doc, "C", "Regulatory Compliance Matrix",
        "To be completed — full IRDAI + DPDPA + PMLA compliance matrix for every data source.",
        ["C.1 Matrix — Data Source × Regulatory Basis × Consent Type × Risk Level","C.2 Mandatory Step 1 Consent Screen Additions","C.3 Anti-Redlining Rules and Quarterly Monitoring Checklist","C.4 Data Processing Agreement Requirements by Vendor","C.5 Data Retention Rules by Source Type"])

    add_appendix(doc, "D", "Enriched STP Payload Schema",
        "To be completed — full TypeScript interface for the STP payload incorporating all alternate data sources.",
        ["D.1 Section A — Identity & Fraud Signals","D.2 Section B — Financial Profile","D.3 Section C — Employment & Occupation","D.4 Section D — Medical & Health","D.5 Section E — Geographic & Environmental Risk","D.6 Section F — Fraud & IIB Signals","D.7 Section G — Nominee Signals","D.8 Section H — Term Insurance Specific"])

    # Footer
    add_page_break(doc)
    para = doc.add_paragraph()
    run  = para.add_run(
        "Document version 1.0 — July 2026  |  Owner: Platform Architecture Team  |  "
        "Review cycle: Quarterly  |  Next review: October 2026"
    )
    run.font.name   = BODY_FONT
    run.font.size   = Pt(9)
    run.font.italic = True
    run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
    para.alignment  = WD_ALIGN_PARAGRAPH.CENTER

    out = r"e:\Insuretech\india-health-platform\docs\ALTERNATE_DATA_SOURCES_REPORT.docx"
    doc.save(out)
    print(f"Saved: {out}")


if __name__ == "__main__":
    main()
